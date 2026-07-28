import os
from typing import List, Dict, Any
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

def read_word_file(file_path: str) -> str:
    """
    Reads a Word (.docx) file and returns its content formatted as markdown.
    Preserves document structure including headings, lists, and tables in their correct order.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)
    markdown_parts = []
    
    # Track list item states to format them properly with blank lines around lists
    in_list = False

    # Iterate through body elements in order to preserve layouts
    for child in doc.element.body:
        if child.tag.endswith('p'):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if not text:
                continue
                
            style_name = p.style.name if p.style else ""
            
            # Identify headings and titles
            if style_name == 'Title':
                if in_list:
                    in_list = False
                    markdown_parts.append("")
                markdown_parts.append(f"# {text}\n")
            elif style_name.startswith('Heading'):
                if in_list:
                    in_list = False
                    markdown_parts.append("")
                # Extract level
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                markdown_parts.append(f"{'#' * level} {text}\n")
            
            # Identify list items
            elif 'List Bullet' in style_name:
                in_list = True
                markdown_parts.append(f"- {text}")
            elif 'List Number' in style_name:
                in_list = True
                markdown_parts.append(f"1. {text}")
            
            # Normal paragraphs
            else:
                if in_list:
                    in_list = False
                    markdown_parts.append("")
                markdown_parts.append(f"{text}\n")
                
        elif child.tag.endswith('tbl'):
            if in_list:
                in_list = False
                markdown_parts.append("")
                
            t = Table(child, doc)
            table_markdown = []
            
            # Determine max columns in the table
            max_cols = 0
            rows_data = []
            for row in t.rows:
                row_cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                # Filter merged cell duplication (python-docx merges cells by repeating references)
                rows_data.append(row_cells)
                max_cols = max(max_cols, len(row_cells))
                
            if max_cols > 0:
                # Format as markdown table
                # Header row
                header_row = rows_data[0] if rows_data else []
                # Fill missing cols if any
                header_row += [""] * (max_cols - len(header_row))
                table_markdown.append("| " + " | ".join(header_row) + " |")
                
                # Separator row
                table_markdown.append("| " + " | ".join(["---"] * max_cols) + " |")
                
                # Data rows
                for r_data in rows_data[1:]:
                    r_data += [""] * (max_cols - len(r_data))
                    table_markdown.append("| " + " | ".join(r_data) + " |")
                
                markdown_parts.append("\n".join(table_markdown) + "\n")
                
    return "\n".join(markdown_parts)


def write_word_file(file_path: str, content: List[Dict[str, Any]], title: str = None) -> str:
    """
    Creates a new Word (.docx) file (or overwrites an existing one) with structured content.
    Content structure:
    List of blocks:
    [
        {"type": "heading", "level": 1, "text": "Introduction"},
        {"type": "paragraph", "text": "This is a normal paragraph."},
        {"type": "list_item", "style": "List Bullet", "text": "Point 1"},
        {"type": "table", "headers": ["Name", "Value"], "rows": [["Item A", "100"], ["Item B", "200"]]}
    ]
    """
    doc = Document()
    
    if title:
        doc.add_heading(title, 0)
        
    for block in content:
        b_type = block.get("type", "paragraph").lower()
        text = block.get("text", "")
        
        if b_type == "heading":
            level = int(block.get("level", 1))
            doc.add_heading(text, level=level)
            
        elif b_type == "paragraph":
            doc.add_paragraph(text)
            
        elif b_type == "list_item":
            style = block.get("style", "List Bullet")
            # Ensure valid list style name
            if "bullet" in style.lower():
                style_name = "List Bullet"
            elif "number" in style.lower():
                style_name = "List Number"
            else:
                style_name = "List Bullet"
            doc.add_paragraph(text, style=style_name)
            
        elif b_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            
            num_rows = len(rows)
            if headers:
                num_rows += 1
            
            num_cols = max(len(headers), max((len(r) for r in rows), default=0))
            if num_cols == 0:
                continue
                
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            curr_row = 0
            if headers:
                hdr_cells = table.rows[curr_row].cells
                for col_idx, header_val in enumerate(headers):
                    hdr_cells[col_idx].text = str(header_val)
                curr_row += 1
                
            for r_data in rows:
                row_cells = table.rows[curr_row].cells
                for col_idx, val in enumerate(r_data):
                    if col_idx < num_cols:
                        row_cells[col_idx].text = str(val)
                curr_row += 1
                
    # Create directory if it doesn't exist
    dir_name = os.path.dirname(file_path)
    if dir_name and not os.path.exists(dir_name):
        os.makedirs(dir_name)
        
    doc.save(file_path)
    return f"Word document successfully written to {file_path}"
